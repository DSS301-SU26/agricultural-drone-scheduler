import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system(sys.executable + " -m pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add a Title
title = doc.add_heading('D2 - MODEL REPORT', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Hệ thống Hỗ trợ Ra quyết định Điều phối Drone Nông nghiệp', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Môn học: ').bold = True
p.add_run('Decision Support Systems for Business Intelligence (DSS301)\n')
p.add_run('Học kỳ: ').bold = True
p.add_run('Summer 2026\n')
p.add_run('Deliverable: ').bold = True
p.add_run('D2 - Model Component')

doc.add_heading('Mục lục', level=1)
toc = [
    "1. Tổng quan bài toán (Decision Problem)",
    "2. Phương pháp luận & Quyết định loại mô hình (Model Selection)",
    "3. Chuẩn bị và Tiền xử lý dữ liệu (Data Preprocessing)",
    "4. Cấu trúc các mô hình học máy (Model Implementation)",
    "   4.1. Mô hình chính: Random Forest",
    "   4.2. Các mô hình so sánh: Logistic Regression & Decision Tree",
    "5. Quá trình huấn luyện & Log Hyperparameter",
    "6. Đánh giá & So sánh mô hình (Evaluation & Model Comparison)",
    "7. AI Audit Log"
]
for item in toc:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 1. Tổng quan bài toán
doc.add_heading('1. Tổng quan bài toán (Decision Problem)', level=1)
doc.add_paragraph('Báo cáo D2 này trình bày chi tiết về thành phần Mô hình (Model Component) của Hệ thống DSS Hỗ trợ Ra quyết định lập lịch bay và vận hành Drone nông nghiệp (Agricultural Drone Scheduler). Mục tiêu cốt lõi của mô hình là xử lý dữ liệu khí tượng (thời tiết) và tình trạng mùa màng để tự động hóa việc đưa ra các khuyến nghị vận hành (decision_action) cho người điều khiển (UAV Operator).')
doc.add_paragraph('Các quyết định này giúp tối ưu hóa hiệu quả phun thuốc/tưới tiêu, đảm bảo an toàn thiết bị (ngăn chặn rơi rớt do gió bão) và giảm thiểu các rủi ro vận hành (như thuốc bị rửa trôi do mưa).')

# 2. Quyết định loại mô hình
doc.add_heading('2. Phương pháp luận & Quyết định loại mô hình (Model Selection)', level=1)
doc.add_paragraph('Biến mục tiêu của bài toán là decision_action với 4 trạng thái rời rạc:')
p = doc.add_paragraph()
p.add_run('• TAKE_OFF: Điều kiện lý tưởng để cất cánh.\n')
p.add_run('• DELAY_FLIGHT: Hoãn chuyến bay do điều kiện chưa phù hợp (nắng gắt, xác suất mưa trung bình).\n')
p.add_run('• LOCK_SPRAY: Khóa lệnh phun và không cất cánh do gió giật mạnh, nguy hiểm cho thiết bị.\n')
p.add_run('• RETURN_TO_CHARGING: Gọi drone về trạm khẩn cấp do có mưa dông.')
doc.add_paragraph('Vì biến mục tiêu là dữ liệu phân loại nhiều lớp (Multi-class), nhóm quyết định sử dụng các thuật toán Machine Learning Classification thay vì Regression. Dù hệ thống có thể xây dựng bằng Rule-based (Tập luật if-else), hướng tiếp cận học máy giúp hệ thống tự phát hiện ra các tương tác đa biến phức tạp giữa thời tiết và môi trường mà con người khó liệt kê hết, đồng thời dễ dàng scale khi áp dụng cho dữ liệu lớn (Big Data) sau này.')

# 3. Tiền xử lý dữ liệu
doc.add_heading('3. Chuẩn bị và Tiền xử lý dữ liệu (Data Preprocessing)', level=1)
doc.add_paragraph('Dữ liệu được thu thập từ file final_training_data.csv. Sau khi làm sạch các bản ghi trùng lặp (dựa trên location_name và timestamp), nhóm đã thiết lập một Pipeline tiền xử lý:')
p = doc.add_paragraph()
p.add_run('• Tách tập Train/Test (75/25): ').bold = True
p.add_run('Nhóm sử dụng kỹ thuật GroupShuffleSplit theo timestamp để đảm bảo mô hình không bị rò rỉ dữ liệu (Data Leakage) giữa các địa điểm cùng chung một hình thái thời tiết.\n')
p.add_run('• Xử lý giá trị thiếu (Missing Values): ').bold = True
p.add_run('Sử dụng SimpleImputer(strategy="median") cho tất cả các mô hình.\n')
p.add_run('• Chuẩn hóa dữ liệu (Scaling): ').bold = True
p.add_run('Sử dụng StandardScaler() riêng cho mô hình tuyến tính (Logistic Regression) để giúp thuật toán hội tụ. Tree-based models không cần bước này.')

# 4. Xây dựng mô hình
doc.add_heading('4. Cấu trúc các mô hình học máy (Model Implementation)', level=1)
doc.add_paragraph('Nhóm triển khai thử nghiệm 3 mô hình học máy và 1 mô hình cơ sở (Dummy Baseline) để so sánh chéo.')

doc.add_heading('4.1. Mô hình chính: Random Forest Classifier', level=2)
doc.add_paragraph('Random Forest được chọn làm ứng viên mô hình chính vì đây là một thuật toán học tập tập hợp (Ensemble). Nó tạo ra hàng trăm cây quyết định và bầu chọn kết quả cuối cùng. Đặc tính này giúp thuật toán xử lý cực tốt các biến phi tuyến tính (ví dụ: nhiệt độ cao nhưng độ ẩm lại thấp) và cực kỳ bền bỉ với dữ liệu ngoại lai (outliers). Hơn nữa, nó cung cấp khả năng Explainable AI (XAI) thông qua Feature Importance.')

doc.add_heading('4.2. Các mô hình so sánh: Logistic Regression & Decision Tree', level=2)
p = doc.add_paragraph()
p.add_run('• Logistic Regression: ').bold = True
p.add_run('Mô hình phân loại tuyến tính kinh điển, đóng vai trò làm baseline so sánh về tốc độ và tính hội tụ. Thuật toán này đòi hỏi tiền xử lý Scaling chặt chẽ.\n')
p.add_run('• Decision Tree: ').bold = True
p.add_run('Thuật toán cơ sở của Random Forest. Việc thêm Decision Tree giúp nhóm đánh giá xem liệu một cây quyết định đơn lẻ có đủ sức giải quyết bài toán hay không, hay bắt buộc phải dùng sức mạnh của Ensemble.')

# 5. Huấn luyện và Hyperparameters
doc.add_heading('5. Quá trình huấn luyện & Log Hyperparameter', level=1)
doc.add_paragraph('Các mô hình được huấn luyện đồng thời thông qua danh sách ứng viên (model_candidates) trên cùng một tập X_train. Dưới đây là bảng thông số Hyperparameter đã được fix lại nhằm đảm bảo khả năng tái lập (reproducibility):')

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Hyperparameters'

data = [
    ('Random Forest', 'n_estimators=250, max_depth=10, min_samples_leaf=4, class_weight="balanced_subsample", random_state=42'),
    ('Logistic Regression', 'max_iter=2000, class_weight="balanced", random_state=42'),
    ('Decision Tree', 'max_depth=6, min_samples_leaf=8, class_weight="balanced", random_state=42')
]
for m, h in data:
    row_cells = table.add_row().cells
    row_cells[0].text = m
    row_cells[1].text = h

# 6. Đánh giá và So sánh
doc.add_heading('6. Đánh giá & So sánh mô hình (Evaluation & Model Comparison)', level=1)
doc.add_paragraph('Sau quá trình Prediction trên tập X_test, các mô hình được so sánh dựa trên các Metrics: Accuracy, Precision, Recall, và đặc biệt là Macro F1. Macro F1 được nhóm ưu tiên nhất vì nó tính trung bình F1-score của tất cả các nhãn (bất kể số lượng mẫu), giúp mô hình không bị thiên vị cho nhãn đa số (TAKE_OFF) mà bỏ sót các rủi ro hiếm gặp (RETURN_TO_CHARGING).')

doc.add_heading('Kết quả so sánh & Giải thích lựa chọn cuối cùng', level=2)
p = doc.add_paragraph()
p.add_run('• Decision Tree: ').bold = True
p.add_run('Mặc dù học nhanh và có thể giải thích dễ dàng thông qua sơ đồ cây, mô hình này dễ bị Overfitting và đạt F1-score thấp hơn các thuật toán khác trên tập test.\n')
p.add_run('• Logistic Regression: ').bold = True
p.add_run('Chạy ổn định, tuy nhiên do ranh giới quyết định (Decision Boundary) của thời tiết khá phức tạp (phi tuyến tính), Logistic Regression bị hạn chế về độ chính xác và số lượng False Positive khá lớn.\n')
p.add_run('• Random Forest: ').bold = True
p.add_run('Đạt hiệu suất (Accuracy & Macro F1) cao nhất trong các thuật toán. Việc kết hợp tham số class_weight="balanced_subsample" giúp thuật toán học rất tốt các trường hợp hiếm như LOCK_SPRAY, đảm bảo rủi ro phần cứng của Drone được phát hiện kịp thời.')
doc.add_paragraph('=> KẾT LUẬN: Random Forest chính thức được lựa chọn làm mô hình cuối cùng (Final Model). Mô hình này sau đó được retrain trên toàn bộ Dataset và export thành file drone_decision_model.joblib để sẵn sàng tích hợp vào Dashboard.')

doc.add_page_break()

# 7. AI Audit Log
doc.add_heading('7. AI Audit Log', level=1)
doc.add_paragraph('Dưới đây là nhật ký tương tác với công cụ AI thể hiện tư duy phản biện (Critical Thinking) và kiểm chứng của nhóm trong Phase 3 (Model Component).')

table_ai = doc.add_table(rows=1, cols=4)
table_ai.style = 'Table Grid'
ai_hdr = table_ai.rows[0].cells
ai_hdr[0].text = 'STT'
ai_hdr[1].text = 'DTC + Phase'
ai_hdr[2].text = 'Prompt (rút gọn)'
ai_hdr[3].text = 'AI Output + Decision & Refinement'

ai_data = [
    [
        '1', 'Decomposition\n(P3 — Model Selection)',
        '"Nên dùng thuật toán nào cho hệ thống dự báo lịch bay Drone: Logistic Regression hay Neural Network?"',
        '- AI Output: Khuyên dùng Neural Network để tối đa hóa độ chính xác.\n- Decision: Bác bỏ. Neural Network là dạng hộp đen (Black-box), làm mất tính giải thích (Explainable AI) rất cần cho hệ thống ra quyết định DSS. Nhóm chọn Random Forest vì mạnh mẽ mà vẫn giải thích được Feature Importance.'
    ],
    [
        '2', 'Pattern Recognition\n(P3 — Data Imbalance)',
        '"Gợi ý cách xử lý dữ liệu bị mất cân bằng khi nhãn TAKE_OFF chiếm tỷ trọng lớn"',
        '- AI Output: Đề xuất SMOTE, Undersampling, class_weight="balanced", và gộp nhãn.\n- Decision: Nhóm bác bỏ SMOTE vì nó có thể sinh ra các bản ghi thời tiết giả (hallucination) không hợp logic. Nhóm quyết định chỉ sử dụng tham số thuật toán class_weight="balanced_subsample" cho Random Forest.'
    ],
    [
        '3', 'Process / Algorithm\n(P3 — Validation)',
        '"Nên chia Train/Test bằng train_test_split ngẫu nhiên không?"',
        '- AI Output: Đồng ý và cho sẵn đoạn code train_test_split.\n- Critical Thinking: Nhóm nhận ra rủi ro Data Leakage nghiêm trọng (thời tiết ở các location gần nhau tại cùng một giờ sẽ giống hệt nhau). Nhóm tự quyết định đổi sang GroupShuffleSplit theo timestamp để mô phỏng dữ liệu tương lai tốt hơn.'
    ],
    [
        '4', 'Abstraction\n(P3 — Preprocessing)',
        '"Viết code chuẩn hóa dữ liệu (Scaler) cho Pipeline."',
        '- AI Output: Sinh code áp dụng StandardScaler cho TẤT CẢ các mô hình.\n- Decision: Tinh chỉnh lại. Mô hình Tree-based (Decision Tree, Random Forest) không cần thiết phải chuẩn hóa. Nhóm tự tách thành 2 Pipeline song song cho Tree và Linear model.'
    ],
    [
        '5', 'Decomposition\n(P3 — Evaluation Metrics)',
        '"Dùng Accuracy để chấm điểm Random Forest là tốt nhất đúng không?"',
        '- AI Output: Xác nhận và sinh code tính Accuracy_score.\n- Rationale: Trong dự báo rủi ro an toàn bay, nhãn nguy hiểm rất hiếm (Imbalanced). Nếu chỉ dùng Accuracy, mô hình sẽ phớt lờ nhãn hiếm. Nhóm bổ sung việc đánh giá bằng Macro F1-score và ưu tiên Metric này.'
    ]
]

for row_data in ai_data:
    row_cells = table_ai.add_row().cells
    for i, text in enumerate(row_data):
        row_cells[i].text = text

doc.save(r'd:\IdeaProjects\DSS301\agricultural-drone-scheduler\submission\D2_Model_Report.docx')
