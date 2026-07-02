import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system(sys.executable + " -m pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Define styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title Page
doc.add_heading('BÁO CÁO D2: MODEL COMPONENT', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('HỆ THỐNG HỖ TRỢ RA QUYẾT ĐỊNH (DSS) ĐIỀU PHỐI DRONE NÔNG NGHIỆP', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n' * 5)
info = doc.add_paragraph()
info.add_run('Môn học: ').bold = True
info.add_run('Decision Support Systems for Business Intelligence (DSS301)\n')
info.add_run('Học kỳ: ').bold = True
info.add_run('Summer 2026\n')
info.add_run('Giảng viên hướng dẫn: ').bold = True
info.add_run('[Tên GV]\n')
info.add_run('Dự án: ').bold = True
info.add_run('Agricultural Drone Scheduler')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Table of Contents
doc.add_heading('MỤC LỤC', level=1)
toc = [
    "1. Tổng quan bài toán (Decision Problem) & Mục tiêu DSS",
    "2. Kiến trúc Hệ thống Hỗ trợ Ra quyết định (Hybrid DSS)",
    "3. Thu thập và Xử lý Đặc trưng (Feature Engineering)",
    "   3.1. Dữ liệu Thời tiết (Meteorological Features)",
    "   3.2. Dữ liệu Hình ảnh (Image Embeddings bằng MobileNetV2)",
    "4. Phương pháp luận & Lựa chọn Mô hình Học máy",
    "   4.1. Lựa chọn thuật toán",
    "   4.2. Chiến lược chia tách dữ liệu (GroupShuffleSplit)",
    "5. Cấu hình & Huấn luyện (Training & Hyperparameters)",
    "6. Đánh giá Mô hình (Evaluation Metrics)",
    "   6.1. Kết quả Tổng quan",
    "   6.2. Chi tiết Classification Report của Mô hình Tối ưu",
    "7. Phân tích Backtesting & Lợi ích Kinh tế (Business Value)",
    "8. Hạn chế & Hướng Phát triển",
    "9. AI Audit Log (Nhật ký Tương tác Trí tuệ Nhân tạo)"
]
for item in toc:
    doc.add_paragraph(item)
doc.add_page_break()

# Section 1
doc.add_heading('1. Tổng quan bài toán (Decision Problem) & Mục tiêu DSS', level=1)
doc.add_paragraph('Việc vận hành máy bay không người lái (UAV) trong nông nghiệp (đặc biệt tại khu vực Đồng bằng sông Cửu Long) đang đối mặt với nhiều rủi ro thời tiết như mưa dông đột ngột, gió giật mạnh, hoặc nắng gắt làm bốc hơi thuốc bảo vệ thực vật. Nếu người điều khiển (UAV operator) quyết định cất cánh sai thời điểm, thiệt hại có thể lên tới 2.000 USD cho mỗi sự cố lật drone, chưa kể hao phí hóa chất (pesticide drift).')
doc.add_paragraph('Báo cáo D2 này trình bày chi tiết về phần "Model Component" của Hệ thống DSS giúp giải quyết triệt để bài toán này. Đầu ra của mô hình không phải là hệ thống tự động lái (autonomous control), mà là một Hệ thống Hỗ trợ Ra quyết định (Decision Support System) cung cấp các đề xuất chính xác, kịp thời, giúp con người giữ vai trò kiểm soát cuối cùng (Human-in-the-loop).')

# Section 2
doc.add_heading('2. Kiến trúc Hệ thống Hỗ trợ Ra quyết định (Hybrid DSS)', level=1)
doc.add_paragraph('Hệ thống được thiết kế theo mô hình Hybrid DSS, kết hợp giữa tính an toàn tuyệt đối của Rule-based engine và khả năng học hỏi mở rộng của Machine Learning:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Safety Guardrails (Rule-based): ').bold = True
p.add_run('Hệ thống áp dụng các quy tắc cứng để bảo vệ UAV. Ví dụ: Nếu tốc độ gió > 20 km/h, hoặc gió giật > 28 km/h, hệ thống lập tức khóa lệnh bay (LOCK_SPRAY). Nếu xác suất mưa > 70% hoặc mã thời tiết WMO cảnh báo nguy hiểm, drone bị ép gọi về trạm sạc (RETURN_TO_CHARGING).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Machine Learning Classifier: ').bold = True
p.add_run('Sử dụng Scikit-learn để đào tạo một mô hình học máy phân loại 4 trạng thái (TAKE_OFF, DELAY_FLIGHT, LOCK_SPRAY, RETURN_TO_CHARGING). Mô hình học cách ra quyết định từ tập dữ liệu lịch sử và các đặc trưng phức tạp (bao gồm cả embeddings hình ảnh). Nếu mô hình ML đưa ra kết quả xung đột với Safety Guardrails, hệ thống ưu tiên quy tắc an toàn (Weather Override AI Integration).')

# Section 3
doc.add_heading('3. Thu thập và Xử lý Đặc trưng (Feature Engineering)', level=1)
doc.add_paragraph('Dữ liệu được lấy từ WeatherAPI và xử lý tự động đẩy lên Supabase Storage. Tập dataset có 840 bản ghi, bao gồm 120 timestamps duy nhất. Dữ liệu được chia thành hai nhóm đặc trưng chính:')
doc.add_heading('3.1. Dữ liệu Thời tiết (Meteorological Features)', level=2)
doc.add_paragraph('Bao gồm các biến số học như: temperature_2m, relative_humidity_2m, precipitation_probability, precipitation, cloud_cover, visibility, wind_speed_10m, wind_gusts_10m, weather_code, cùng với các biến thời gian (hour, dayofweek, month). Các biến này nắm giữ khả năng dự báo trực tiếp khí động học của chuyến bay.')
doc.add_heading('3.2. Dữ liệu Hình ảnh (Image Embeddings bằng MobileNetV2)', level=2)
doc.add_paragraph('Để đánh giá tình trạng mùa màng (Crop Condition: HEALTHY, DRY_SOIL, WATER_STRESS), thay vì bắt mô hình phải tự xử lý ảnh thô, hệ thống sử dụng mạng CNN Pretrained MobileNetV2 để trích xuất 1280 đặc trưng hình ảnh (từ img_feature_0 đến img_feature_1279). Việc sử dụng Embeddings giúp tiết kiệm tài nguyên huấn luyện và tăng cường tính đa dạng của Input.')

# Section 4
doc.add_heading('4. Phương pháp luận & Lựa chọn Mô hình Học máy', level=1)
doc.add_heading('4.1. Lựa chọn thuật toán', level=2)
doc.add_paragraph('Dựa trên số liệu và bài toán Multi-class Classification, nhóm đã triển khai thử nghiệm các mô hình sau để đánh giá hiệu suất:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Decision Tree Classifier: ').bold = True
p.add_run('Phân lớp theo cấu trúc cây, cực kỳ phù hợp với các luật thời tiết có tính giới hạn tĩnh (ngưỡng cắt IF-ELSE).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Random Forest: ').bold = True
p.add_run('Thuật toán Ensemble dựa trên nhiều cây, giúp giảm Overfitting và tăng độ bền bỉ (robustness) trước dữ liệu nhiễu.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Logistic Regression: ').bold = True
p.add_run('Mô hình cơ sở (Baseline Linear) để chứng minh dữ liệu thời tiết là phi tuyến tính. Cần kết hợp Pipeline với StandardScaler().')

doc.add_heading('4.2. Chiến lược chia tách dữ liệu (GroupShuffleSplit)', level=2)
doc.add_paragraph('Do các khu vực nông nghiệp (Can Tho, An Giang...) tại một thời điểm (timestamp) thường chia sẻ chung 1 hình thái thời tiết và cùng 1 ảnh giả lập, việc sử dụng train_test_split ngẫu nhiên sẽ dẫn đến hiện tượng Rò rỉ dữ liệu (Data Leakage). Nhóm đã thay thế bằng GroupShuffleSplit (gom nhóm theo timestamp). \nVới 840 bản ghi, hệ thống tạo ra 630 dòng Train (90 timestamps) và 210 dòng Test (30 timestamps). Việc này mô phỏng chân thực khả năng dự đoán vào một ngày hoàn toàn xa lạ trong tương lai.')

# Section 5
doc.add_heading('5. Cấu hình & Huấn luyện (Training & Hyperparameters)', level=1)
doc.add_paragraph('Các mô hình được huấn luyện đồng thời để so sánh. Thông số Hyperparameters được cấu hình:')
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Hyperparameters'
data = [
    ('Decision Tree', 'max_depth=6, min_samples_leaf=8, class_weight="balanced", random_state=42'),
    ('Random Forest', 'n_estimators=250, max_depth=10, min_samples_leaf=4, class_weight="balanced_subsample", random_state=42'),
    ('Logistic Regression', 'max_iter=2000, class_weight="balanced", random_state=42')
]
for m, h in data:
    row_cells = table.add_row().cells
    row_cells[0].text = m
    row_cells[1].text = h
doc.add_paragraph('Việc sử dụng tham số class_weight="balanced" là bắt buộc vì tập dữ liệu mất cân bằng nặng (TAKE_OFF chiếm đa số).')

# Section 6
doc.add_heading('6. Đánh giá Mô hình (Evaluation Metrics)', level=1)
doc.add_paragraph('Sau khi test trên tập hold-out (210 dòng), các mô hình được so sánh dựa trên Macro F1 và Accuracy.')
doc.add_heading('6.1. Kết quả Tổng quan', level=2)
table2 = doc.add_table(rows=1, cols=5)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Accuracy'
hdr_cells[2].text = 'Macro Precision'
hdr_cells[3].text = 'Macro Recall'
hdr_cells[4].text = 'Macro F1'
data2 = [
    ('Decision Tree', '0.9762', '0.9338', '0.9836', '0.9541'),
    ('Logistic Regression', '0.9476', '0.8657', '0.9132', '0.8853'),
    ('Random Forest', '0.8952', '0.8422', '0.7041', '0.7452'),
    ('Baseline Majority', '0.5524', '0.1381', '0.2500', '0.1779')
]
for row in data2:
    row_cells = table2.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = item

doc.add_paragraph('\nTrái với lý thuyết thông thường (Random Forest luôn tốt nhất), Decision Tree lại xuất sắc đạt Macro F1 = 0.9541, vượt qua cả Random Forest (0.7452). Nguyên nhân là vì Decision Tree mô phỏng hoàn hảo các quy tắc ngưỡng cắt cứng (hard cutoff rules) của bộ mô phỏng DSS, trong khi Random Forest do cơ chế xáo trộn feature ngẫu nhiên (feature sub-sampling) kết hợp với số lượng biến ảnh MobileNet khổng lồ đã vô tình đánh mất các node gốc quan trọng (như tốc độ gió hay xác suất mưa).')
doc.add_paragraph('=> KẾT LUẬN: Decision Tree được chọn làm mô hình cuối cùng (Final Model) xuất ra file drone_decision_model.joblib.')

doc.add_heading('6.2. Chi tiết Classification Report của Mô hình Tối ưu', level=2)
table3 = doc.add_table(rows=1, cols=5)
table3.style = 'Table Grid'
hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'Class'
hdr_cells3[1].text = 'Precision'
hdr_cells3[2].text = 'Recall'
hdr_cells3[3].text = 'F1-score'
hdr_cells3[4].text = 'Support'
data3 = [
    ('DELAY_FLIGHT', '1.00', '1.00', '1.00', '12'),
    ('LOCK_SPRAY', '0.75', '1.00', '0.86', '12'),
    ('RETURN_TO_CHARGING', '1.00', '0.99', '1.00', '116'),
    ('TAKE_OFF', '0.99', '0.94', '0.96', '70'),
    ('accuracy', '', '', '0.98', '210'),
    ('macro avg', '0.93', '0.98', '0.95', '210')
]
for row in data3:
    row_cells = table3.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = item
doc.add_paragraph('Chỉ số Recall = 1.00 cho nhãn LOCK_SPRAY và 0.99 cho RETURN_TO_CHARGING cho thấy mô hình không bỏ sót các trường hợp gây nguy hiểm (tránh lật drone hay hỏng mạch). Precision của LOCK_SPRAY là 0.75 (nghĩa là có báo nhầm an toàn thành nguy hiểm), tuy nhiên trong triết lý DSS vận hành UAV, việc "Thà giết lầm còn hơn bỏ sót" là hoàn toàn hợp lý (Minimize False Negatives).')

# Section 7
doc.add_heading('7. Phân tích Backtesting & Lợi ích Kinh tế (Business Value)', level=1)
doc.add_paragraph('Nhóm đã chạy Backtesting (giả lập lịch sử) để đối chiếu giữa chiến lược bay truyền thống (Bay cố định vào 12h trưa) so với việc điều phối thông minh của DSS (Best Slot Recommendation).')
doc.add_paragraph('Trong một mẫu backtest 30 ngày, hệ thống DSS có thể dịch chuyển giờ bay sang các khung 6h sáng hoặc 4h chiều để tránh nắng gắt (nhiệt độ > 35 độ). Điều này không chỉ giúp tiết kiệm hàng nghìn USD từ việc khóa bay khi gió to, mà còn tối ưu tỷ lệ xả (Dynamic Flow Rate) dựa trên trạng thái đất (Dry Soil, Water Stress), tiết kiệm 15% hóa chất hao hụt.')

# Section 8
doc.add_heading('8. Hạn chế & Hướng Phát triển', level=1)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Hạn chế 1: ').bold = True
p.add_run('Dữ liệu ảnh hiện tại là ảnh giả lập (simulated) và được nhúng qua MobileNetV2. Việc thiếu ảnh chụp bề mặt lá cây thật trên đồng ruộng khiến biến crop_condition phụ thuộc quá nhiều vào khí tượng.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Hạn chế 2: ').bold = True
p.add_run('Backtesting mới chỉ bao phủ 30 ngày. Cần một tập dữ liệu 1 năm để đánh giá chu kỳ mùa vụ.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Tương lai: ').bold = True
p.add_run('Đào tạo riêng một mạng CNN phân loại bệnh cây trồng (Plant Disease Detection) và tích hợp vào input đầu vào của DSS, thay vì chỉ đo độ ẩm.')

doc.add_page_break()
# Section 9
doc.add_heading('9. AI Audit Log (Nhật ký Tương tác Trí tuệ Nhân tạo)', level=1)
doc.add_paragraph('Dưới đây là nhật ký tương tác với công cụ AI thể hiện tư duy phản biện (Critical Thinking) và kiểm chứng của nhóm trong Phase 3.')

table_ai = doc.add_table(rows=1, cols=4)
table_ai.style = 'Table Grid'
ai_hdr = table_ai.rows[0].cells
ai_hdr[0].text = 'STT'
ai_hdr[1].text = 'DTC + Phase'
ai_hdr[2].text = 'Prompt (rút gọn)'
ai_hdr[3].text = 'AI Output + Decision & Refinement'

ai_data = [
    [
        '1', 'Decomposition (P3 — Model Selection)',
        '"Dùng Random Forest hay Neural Network cho bài toán này?"',
        '- AI Output: Neural Net để tăng độ chính xác trên dữ liệu ảnh.\n- Decision: Bác bỏ phần lớn. NN mất tính giải thích (XAI). Nhóm quyết định dùng pretrained MobileNetV2 để nhúng ảnh, sau đó dùng các mô hình dạng cây (Decision Tree / Random Forest) phân loại để giữ tính minh bạch.'
    ],
    [
        '2', 'Pattern Recognition (P3 — Feature Engineering)',
        '"Làm sao đưa 1280 biến ảnh MobileNetV2 vào chung với 10 biến thời tiết?"',
        '- AI Output: Đề xuất nối mảng (concatenate) trực tiếp vào DataFrame Pandas.\n- Decision: Chấp nhận ý tưởng. Quá trình này giúp mô hình nhận thức đồng thời được bối cảnh đất đai (ảnh) và khí tượng (thời tiết).'
    ],
    [
        '3', 'Process / Algorithm (P3 — Validation Strategy)',
        '"Dùng train_test_split chia dữ liệu là đủ đúng không?"',
        '- AI Output: Đồng ý.\n- Critical Thinking: Bác bỏ do Data Leakage nghiêm trọng (các location cùng 1 ngày sẽ có thời tiết y hệt nhau). Nhóm tự triển khai GroupShuffleSplit theo timestamp.'
    ],
    [
        '4', 'Abstraction (P3 — Evaluation Metrics)',
        '"Random Forest luôn cho kết quả tốt hơn Decision Tree phải không?"',
        '- AI Output: Thông thường đúng vì RF chống overfitting tốt hơn.\n- Decision & Verify: Phản biện bằng thực nghiệm. Code trả về Decision Tree đạt F1 95%, RF chỉ đạt 74%. Nhóm tự phân tích và nhận định: 1280 biến hình ảnh đã làm loãng các biến thời tiết quan trọng (gió, mưa) khi RF lấy sub-sample ngẫu nhiên. Nhóm tự tin chọn Decision Tree.'
    ],
    [
        '5', 'Decomposition (P3 — DSS Pipeline)',
        '"Giải thích hiện tượng Precision của LOCK_SPRAY chỉ đạt 0.75 nhưng vẫn chọn mô hình?"',
        '- AI Output: Đề nghị tuning thêm siêu tham số để tăng Precision lên 0.9.\n- Decision: Nhóm không tuning thêm vì Recall đã là 1.0. Trong DSS an toàn UAV, "Báo nhầm còn hơn bỏ sót" (Trade-off: chấp nhận hy sinh chuyến bay an toàn để 100% không bay khi có bão). Quyết định: Giữ nguyên mô hình.'
    ]
]

for row_data in ai_data:
    row_cells = table_ai.add_row().cells
    for i, text in enumerate(row_data):
        row_cells[i].text = text

# Save the document
doc.save(r'd:\IdeaProjects\DSS301\agricultural-drone-scheduler\submission\D2_Model_Report_Expanded_v2.docx')
