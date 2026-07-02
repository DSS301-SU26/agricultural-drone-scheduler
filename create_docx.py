import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system(sys.executable + " -m pip install python-docx")
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add a Title
title = doc.add_heading('BÁO CÁO TUẦN 5', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Dự án: ').bold = True
p.add_run('Hệ thống Hỗ trợ Ra quyết định Điều phối Drone Nông nghiệp (Agricultural Drone Scheduler DSS)\n')
p.add_run('Môn học: ').bold = True
p.add_run('Decision Support Systems for Business Intelligence (DSS301)\n')
p.add_run('Học kỳ: ').bold = True
p.add_run('Summer 2026\n')
p.add_run('Trạng thái tài liệu: ').bold = True
p.add_run('Báo cáo Tuần 5')

doc.add_heading('Contents', level=1)
toc = [
    "1. Giới thiệu",
    "2. Quyết định loại mô hình",
    "3. Chuẩn bị dữ liệu",
    "4. Tiền xử lý dữ liệu",
    "5. Mô hình chính và mô hình so sánh",
    "   5.1. Mô hình chính: Random Forest",
    "   5.2. Mô hình so sánh: Logistic Regression & Decision Tree",
    "6. Training trên dataset UAV",
    "7. Log hyperparameter",
    "8. Prediction trên tập test",
    "9. AI Audit Log"
]
for item in toc:
    doc.add_paragraph(item, style='List Bullet')

# 1. Giới thiệu
doc.add_heading('1. Giới thiệu', level=1)
doc.add_paragraph('Trong giai đoạn này, nhóm thực hiện xây dựng mô hình học máy (Machine Learning) hỗ trợ ra quyết định (DSS) cho bài toán lập lịch bay và vận hành Drone nông nghiệp dựa trên dữ liệu khí tượng (thời tiết) và tình trạng mùa màng. Mục tiêu của mô hình là dự đoán quyết định hành động (decision_action) nhằm đánh giá xem thiết bị có nên cất cánh hay không. Đây là bài toán phân loại nhiều nhãn (Multi-class Classification), phù hợp với hướng tiếp cận Machine Learning Classification hơn là Regression.')
doc.add_paragraph('Dữ liệu sử dụng gồm các thuộc tính liên quan đến thời tiết (nhiệt độ, độ ẩm, tốc độ gió, sức gió giật, lượng mưa, xác suất mưa, tầm nhìn), thời gian và các đặc trưng hình ảnh đồng ruộng. Trên cơ sở đó, nhóm tiến hành tiền xử lý dữ liệu, xây dựng mô hình chính, mô hình so sánh, và đánh giá kết quả trên tập kiểm tra.')

# 2. Quyết định loại mô hình
doc.add_heading('2. Quyết định loại mô hình', level=1)
doc.add_paragraph('Sau khi phân tích bản chất của biến mục tiêu decision_action với 4 trạng thái phân loại riêng biệt (TAKE_OFF, DELAY_FLIGHT, LOCK_SPRAY, RETURN_TO_CHARGING), nhóm quyết định lựa chọn ML Classification làm phương pháp chính. Lý do là decision_action là biến phân loại, không phải biến số liên tục. Vì vậy, regression không phù hợp với bài toán này.')
doc.add_paragraph('Dù hệ thống có xây dựng tập luật (rule-based) để gán nhãn dữ liệu chuẩn ban đầu, nhưng việc đào tạo mô hình học máy sẽ giúp hệ thống phân tích và dự đoán (inference) với tốc độ cao trên quy mô dữ liệu lớn, cũng như tự học được các tương tác đa chiều phức tạp giữa môi trường và thời tiết.')

# 3. Chuẩn bị dữ liệu
doc.add_heading('3. Chuẩn bị dữ liệu', level=1)
doc.add_paragraph('Dữ liệu được thu thập và tổng hợp thông qua quy trình tự động lấy thông tin từ WeatherAPI và lưu trữ tại final_training_data.csv. Trong quá trình chuẩn bị, dữ liệu trùng lặp (các bản ghi có location_name và timestamp giống nhau) được làm sạch và chỉ giữ lại bản ghi mới nhất.')
p = doc.add_paragraph('Biến mục tiêu decision_action gồm 4 lớp:\n')
p.add_run('• TAKE_OFF: Điều kiện lý tưởng.\n')
p.add_run('• DELAY_FLIGHT: Hoãn do rủi ro thời tiết (mưa, nắng nóng).\n')
p.add_run('• LOCK_SPRAY: Khóa lệnh bay do gió giật vượt ngưỡng.\n')
p.add_run('• RETURN_TO_CHARGING: Bắt buộc thu hồi do có mưa.')
doc.add_paragraph('Các đặc trưng đầu vào chủ yếu là các biến số học (Numeric variables): temperature_2m, relative_humidity_2m, precipitation_probability, precipitation, cloud_cover, visibility, wind_speed_10m, wind_gusts_10m, weather_code, hour, dayofweek, month (cùng các biến đặc trưng hình ảnh nếu có).')
doc.add_paragraph('Dữ liệu được chia Train/Test theo tỷ lệ 75/25. Để bài toán phản ánh đúng thực tế dự báo thời tiết mới, nhóm áp dụng kỹ thuật GroupShuffleSplit nhóm theo thuộc tính timestamp để ngăn chặn hiện tượng rò rỉ dữ liệu (data leakage) giữa các địa điểm trong cùng một khung giờ.')

# 4. Tiền xử lý dữ liệu
doc.add_heading('4. Tiền xử lý dữ liệu', level=1)
doc.add_paragraph('Do dữ liệu đầu vào là biến số và có thể có giá trị thiếu, nhóm sử dụng pipeline tiền xử lý để bảo đảm quá trình train và test được nhất quán. Cụ thể:')
p = doc.add_paragraph()
p.add_run('• Với mô hình Tree-based (Random Forest, Decision Tree): ').bold = True
p.add_run('Sử dụng SimpleImputer(strategy="median") để điền các giá trị bị thiếu bằng giá trị trung vị.\n')
p.add_run('• Với mô hình tuyến tính (Logistic Regression): ').bold = True
p.add_run('Dữ liệu được bổ sung thêm bước chuẩn hóa thang đo bằng StandardScaler() giúp thuật toán hội tụ tốt hơn.')
doc.add_paragraph('Việc sử dụng ColumnTransformer và Pipeline giúp toàn bộ quá trình tiền xử lý và huấn luyện được đóng gói rõ ràng, tránh lỗi rò rỉ dữ liệu từ tập Test sang Train.')

# 5. Mô hình chính và mô hình so sánh
doc.add_heading('5. Mô hình chính và mô hình so sánh', level=1)
doc.add_paragraph('Nhóm triển khai các mô hình để so sánh hiệu quả:')
doc.add_heading('5.1. Mô hình chính: Random Forest', level=2)
doc.add_paragraph('Random Forest được chọn làm mô hình chính vì đây là một thuật toán ensemble cực kỳ mạnh mẽ, xử lý rất tốt các tương tác phi tuyến tính của dữ liệu thời tiết (ví dụ: nhiệt độ cao nhưng độ ẩm lại thấp). Mô hình ít bị ảnh hưởng bởi giá trị ngoại lai (outliers) và xử lý khá ổn định các class mất cân bằng.')
doc.add_paragraph('Hyperparameter sử dụng: n_estimators = 250, max_depth = 10, class_weight = "balanced_subsample"')

doc.add_heading('5.2. Mô hình so sánh: Logistic Regression & Decision Tree', level=2)
doc.add_paragraph('Hai thuật toán được sử dụng để đối chiếu:')
p = doc.add_paragraph()
p.add_run('• Logistic Regression: ').bold = True
p.add_run('Phù hợp làm baseline cho mô hình tuyến tính, có tốc độ xử lý siêu nhanh nhưng có thể giới hạn khả năng học nếu dữ liệu phi tuyến tính. (max_iter = 2000, class_weight = "balanced").\n')
p.add_run('• Decision Tree: ').bold = True
p.add_run('Mô hình cây quyết định đơn giản, độ phức tạp thấp, làm cơ sở so sánh hiệu năng cho thuật toán Random Forest. (max_depth = 6).')
doc.add_paragraph('Các mô hình được huấn luyện trên cùng một tập train và đánh giá trên cùng tập test để đảm bảo tính công bằng.')

# 6. Training trên dataset UAV
doc.add_heading('6. Training trên dataset UAV', level=1)
doc.add_paragraph('Sau khi dữ liệu được làm sạch và chia train/test (dùng GroupShuffleSplit), nhóm khởi chạy pipeline tự động. Hệ thống sẽ lặp qua danh sách các mô hình và gọi hàm fit() trên X_train.')
p = doc.add_paragraph('Quá trình training cho thấy:\n')
p.add_run('• Random Forest và Decision Tree thường xuyên bám sát thực tế các luật giới hạn phi tuyến tính do chúng chia cắt dữ liệu theo các ngưỡng điều kiện khá tốt.\n')
p.add_run('• Mô hình nào cho ra chỉ số Macro F1 tốt nhất sẽ được lưu thành deployment model (drone_decision_model.joblib) để phục vụ phía server.')

# 7. Log hyperparameter
doc.add_heading('7. Log hyperparameter', level=1)
doc.add_paragraph('Nhóm đã ghi lại các hyperparameter chính được định nghĩa trong mã nguồn nhằm phục vụ việc báo cáo và tái lập kết quả huấn luyện:')

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Hyperparameters'

row_cells = table.add_row().cells
row_cells[0].text = 'Random Forest'
row_cells[1].text = 'n_estimators=250, max_depth=10, min_samples_leaf=4, class_weight="balanced_subsample", random_state=42'

row_cells = table.add_row().cells
row_cells[0].text = 'Logistic Regression'
row_cells[1].text = 'max_iter=2000, class_weight="balanced", random_state=42'

row_cells = table.add_row().cells
row_cells[0].text = 'Decision Tree'
row_cells[1].text = 'max_depth=6, min_samples_leaf=8, class_weight="balanced", random_state=42'

doc.add_paragraph('\nViệc log hyperparameter giúp đảm bảo mô hình có thể kiểm tra lại và làm nền tảng cho việc tinh chỉnh (fine-tuning) sau này.')

# 8. Prediction trên tập test
doc.add_heading('8. Prediction trên tập test', level=1)
doc.add_paragraph('Cả ba mô hình được sử dụng để dự đoán trên X_test. Do tập dữ liệu này bao gồm những khung thời gian mới hoàn toàn (nhờ GroupShuffleSplit), mô hình được đánh giá khắt khe nhất dựa trên:')
p = doc.add_paragraph()
p.add_run('• Accuracy: ').bold = True
p.add_run('Độ chính xác tổng thể.\n')
p.add_run('• Macro Precision / Macro Recall / Macro F1: ').bold = True
p.add_run('Rất quan trọng, để đảm bảo hệ thống không dự đoán thiên lệch về nhãn quá đa số như TAKE_OFF mà bỏ quên các nhãn nguy hiểm như LOCK_SPRAY hay RETURN_TO_CHARGING.')
doc.add_paragraph('Kết quả cuối cùng và Classification Report được tự động sinh ra và ghi đè vào thư mục reports. Căn cứ vào đó, nhóm có thể khẳng định mô hình chính (Random Forest) cho khả năng ổn định và phát hiện lỗi do điều kiện thời tiết khắc nghiệt một cách an toàn nhất.')

# 9. AI Audit Log
doc.add_heading('9. AI Audit Log', level=1)
doc.add_paragraph('Dưới đây là nhật ký tương tác với công cụ AI thể hiện tư duy phản biện (Critical Thinking) và việc kiểm chứng dữ liệu thực tế (Human Delta) chiếu theo yêu cầu chuẩn môn DSS301.')

table_ai = doc.add_table(rows=1, cols=4)
table_ai.style = 'Table Grid'
ai_hdr = table_ai.rows[0].cells
ai_hdr[0].text = 'STT'
ai_hdr[1].text = 'DTC + Phase'
ai_hdr[2].text = 'Prompt (rút gọn)'
ai_hdr[3].text = 'AI Output + Decision & Refinement'

ai_data = [
    [
        '1', 'Decomposition(W2 — Target Users)',
        '"Liệt kê các stakeholders (bên liên quan) cho hệ thống DSS hỗ trợ lập lịch bay UAV nông nghiệp dựa trên thời tiết."',
        '- AI Output: 7 stakeholders (nông dân, phi công UAV, cơ quan quản lý không phận, nhà cung cấp API thời tiết, đội bảo trì, công ty bảo hiểm, kỹ sư nông nghiệp).\n- Decision: Chấp nhận 3 nhóm cốt lõi: farm managers, UAV pilots, và agricultural engineers. Loại bỏ công ty bảo hiểm, đội bảo trì và cơ quan không phận vì nằm ngoài scope của một DSS tối ưu thời tiết.\n- Rationale: Giữ cho hệ thống tập trung trực tiếp vào người ra quyết định vận hành (Decision Makers) thay vì các bên quản lý vĩ mô.'
    ],
    [
        '2', 'Decomposition(W2 — AS-IS / Pain points)',
        '"Phân rã các rủi ro và vấn đề chính (pain points) khi vận hành UAV phun thuốc trong điều kiện thời tiết xấu."',
        '- AI Output: Liệt kê rủi ro: hóa chất bay tạt (drift), hỏng pin do lạnh, rơi rớt drone do gió, và phạt vi phạm quy định bay.\n- Decision: Giữ lại pesticide drift và rơi rớt do gió lớn. Bổ sung thêm wash-off risk (rửa trôi hóa chất do mưa dông cục bộ).\n- Critical Thinking: AI đưa ra "hỏng pin do lạnh", rủi ro này chỉ đúng ở xứ ôn đới, hoàn toàn không phù hợp với khí hậu nhiệt đới tại ĐBSCL (Việt Nam) nên nhóm đã phản biện và loại bỏ.'
    ],
    [
        '3', 'Pattern Recognition(W3 — Data Variables)',
        '"Đề xuất các biến dữ liệu thời tiết (weather features) quan trọng nhất để xây dựng logic đánh giá an toàn cho chuyến bay UAV."',
        '- AI Output: Tốc độ gió, nhiệt độ, độ ẩm, áp suất khí quyển, chỉ số UV, lượng mây, lượng mưa, tầm nhìn.\n- Decision: Chỉ chọn 6 biến cốt lõi: temperature, humidity, wind speed, wind direction, precipitation, visibility. Khái quát hóa và loại bỏ UV, áp suất khí quyển.\n- Rationale: Các yếu tố như chỉ số UV không làm thay đổi khí động học của Drone hay chất lượng phun thuốc, nên không có giá trị quyết định (decision-value) trong dataset.'
    ],
    [
        '4', 'Abstraction(W3 — Decision Outputs)',
        '"Hệ thống DSS này nên đưa ra những quyết định đầu ra (decision outputs) nào cho người điều hành UAV?"',
        '- AI Output: Đề xuất 5 output: Cho phép bay, Thay pin, Đổi loại thuốc phun, Chuyển hướng bay, Trì hoãn chuyến bay.\n- Decision: Trừu tượng hóa (Abstraction) và gộp lại thành 3 decision labels chính thức cho Data Dictionary: flight_suitability, operation_risk, và is_delayed.\n- Rationale: Loại bỏ quyết định "Đổi loại thuốc" và "Thay pin" vì nó thuộc về hệ thống quản lý vật tư (Inventory System), không thuộc phạm vi của Weather DSS.'
    ],
    [
        '5', 'Pattern Recognition(W4 — Risk Register)',
        '"Liệt kê 5 rủi ro/hạn chế (limitations) điển hình khi triển khai mô hình dự báo thời tiết cho UAV nông nghiệp."',
        '- AI Output: 5 rủi ro chung chung: trễ tiến độ dự án, lỗi server, API không phản hồi, thiếu dữ liệu, giao diện khó dùng.\n- Decision: Loại bỏ hoàn toàn các rủi ro IT thông thường. Thay bằng 2 rủi ro đặc thù của DSS: (1) Microclimate deviation (sai lệch vi khí hậu giữa API và thực tế ruộng) và (2) Human-in-the-loop dependency (sự quá phụ thuộc vào máy móc).\n- Hallucination/Bias: AI bị thiên kiến (bias) tự động liệt kê các "System/IT Risks" của môn Software Engineering thay vì tập trung vào "Data/Decision Risks" của môn DSS. Nhóm đã tự hiệu chỉnh.'
    ]
]

for row_data in ai_data:
    row_cells = table_ai.add_row().cells
    for i, text in enumerate(row_data):
        row_cells[i].text = text

doc.save(r'd:\IdeaProjects\DSS301\agricultural-drone-scheduler\submission\BaoCao_Tuan5_DSS301_kem_AILog_v2.docx')
