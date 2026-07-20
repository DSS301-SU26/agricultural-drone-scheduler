from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add a title
title = doc.add_heading('BÁO CÁO ĐÁNH GIÁ MÔ HÌNH HỌC MÁY (EVALUATION REPORT)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add content
content = """
Dự án: DSS301 - Lập lịch bay UAV Nông nghiệp chuyên canh lúa nước
Ngày xuất báo cáo: 17/07/2026

1. TỔNG QUAN TẬP DỮ LIỆU
Nguồn dữ liệu đánh giá: Dữ liệu thời tiết lịch sử kết hợp cấu hình thiết bị bay (Drone) và các giai đoạn sinh trưởng của cây lúa.
Thời gian lấy mẫu: Từ 01/01/2020 đến 08/06/2026
Tổng số lượng mẫu (bản ghi): 185,304 dòng
Số lượng biến nội suy (Features): 20 biến độc lập

Phân bổ nhãn quyết định:
- Nhãn 3 (Trạng thái bình thường - PROCEED): 62,641 mẫu
- Nhãn 0 (Khuyến cáo do gió/thời tiết): 55,990 mẫu
- Nhãn 2 (Điều chỉnh theo môi trường): 55,271 mẫu
- Nhãn 1 (Không an toàn - Rủi ro cao): 11,402 mẫu

(Ghi chú: Dữ liệu phản ánh đúng thực tế khi điều kiện bay bình thường và điều kiện có gió/chướng ngại môi trường chiếm đa số).

2. QUÁ TRÌNH CHIA TẬP DỮ LIỆU
Mô hình sử dụng chiến lược GroupShuffleSplit (nhóm theo mốc thời gian timestamp) để đảm bảo không bị rò rỉ dữ liệu tương lai vào tập huấn luyện (Data Leakage).
- Tập Huấn luyện (Training Set): 138,978 mẫu (~75% dữ liệu)
- Tập Kiểm thử (Testing Set): 46,326 mẫu (~25% dữ liệu)
- Số lượng dữ liệu trùng lặp bị loại bỏ: 0

3. KẾT QUẢ ĐÁNH GIÁ (MODEL METRICS)
Hệ thống đã tiến hành huấn luyện song song và so sánh hiệu suất giữa 2 thuật toán:

Thuật toán XGBoost Classifier:
- Độ chính xác tổng thể (Accuracy): 95.04%
- Macro Precision: 89.04%
- Macro Recall: 83.19%
- Chỉ số F1-Score (Macro): 84.85%
- Chỉ số F1-Score (Weighted): 94.40%

Thuật toán Random Forest Classifier:
- Độ chính xác tổng thể (Accuracy): 89.65%
- Macro Precision: 82.40%
- Macro Recall: 87.59%
- Chỉ số F1-Score (Macro): 82.46%
- Chỉ số F1-Score (Weighted): 90.98%

KẾT LUẬN: Mô hình XGBoost cho ra hiệu suất vượt trội (Độ chính xác đạt hơn 95%) và có khả năng xử lý tốt sự chênh lệch (Imbalance) giữa các nhãn dữ liệu. Do đó, thuật toán này sẽ được ưu tiên làm lõi quyết định chính cho hệ thống sinh lời khuyên bay.

4. VÍ DỤ MINH HỌA (RECOMMENDATION ENGINE TEST)
Hệ thống thử nghiệm bốc ngẫu nhiên một tình huống môi trường trong tập Test (Thời điểm: 06:00 - 01/06/2026 tại trạm Tiền Giang). Mô hình đưa ra phán đoán hoàn toàn chính xác với điểm Flyability Score tuyệt đối (1.0).

Lời khuyên sinh ra tự động từ Model:
"Điều kiện bay an toàn — Gió 1.1 km/h (ngưỡng drone: 29 km/h), nhiệt độ 24.4°C, xác suất mưa chỉ 36%. Lúa giai đoạn Mạ: nên bay cao 2.5-3m và tốc độ nhanh để tránh downwash dập mạ non. Đề xuất lưu lượng xả 15.0 L/ha."
"""

for line in content.strip().split('\n'):
    if line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. '):
        doc.add_heading(line, level=1)
    elif line.startswith('Thuật toán') or line.startswith('Phân bổ') or line.startswith('KẾT LUẬN'):
        doc.add_heading(line, level=2)
    elif line.startswith('Lời khuyên'):
        p = doc.add_paragraph()
        p.add_run(line).italic = True
    else:
        doc.add_paragraph(line)

doc.save("models/evaluation_report.docx")
