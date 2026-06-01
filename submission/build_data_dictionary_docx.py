from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path("data/data_dictionary.md")
OUTPUT = Path("submission/w4_clean_dataset_data_dictionary/data_dictionary.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
HEADER_FILL = "E8EEF5"
HEADER_DARK = "1F4D78"
ROW_ALT_FILL = "F7FAFC"
NOTE_FILL = "F4F6F9"
GRID = "B7C4D3"
INK = "1F2933"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGIN.items():
        elem = tc_mar.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tc_mar.append(elem)
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), GRID)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_inline(paragraph, text, size=11, color=INK, bold=False, italic=False):
    token_pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|https?://\S+)")
    pos = 0
    for match in token_pattern.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos:match.start()]), size=size, color=color, bold=bold, italic=italic)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(size - 0.5, 8), color=DARK_BLUE, bold=bold, italic=italic)
        elif token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, color=color, bold=True, italic=italic)
        else:
            set_run_font(paragraph.add_run(token.rstrip(".,)")), size=size, color=BLUE, bold=bold, italic=italic)
            if token[-1:] in ".,)":
                set_run_font(paragraph.add_run(token[-1]), size=size, color=color, bold=bold, italic=italic)
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), size=size, color=color, bold=bold, italic=italic)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def add_body(doc, text, style=None, before=0, after=6, line=1.25):
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_spacing(paragraph, before=before, after=after, line=line)
    add_inline(paragraph, text)
    return paragraph


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, NOTE_FILL)
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.15)
    add_inline(paragraph, text, size=10.2, color=DARK_BLUE, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_summary_strip(doc):
    metrics = [
        ("504", "RAW ROWS"),
        ("252", "CLEAN ROWS"),
        ("22", "COLUMNS"),
        ("3", "FORECAST DAYS"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [2340, 2340, 2340, 2340]
    set_table_geometry(table, widths)
    set_table_borders(table)
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, HEADER_DARK)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)
        set_run_font(paragraph.add_run(value), size=15, color=WHITE, bold=True)
        label_paragraph = cell.add_paragraph()
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(label_paragraph, before=0, after=0, line=1.0)
        set_run_font(label_paragraph.add_run(label), size=8.2, color="DCE8F3", bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def choose_widths(rows):
    cols = len(rows[0])
    if cols == 2:
        return [2500, 6860]
    if cols == 3:
        return [1650, 3610, 4100]
    if cols == 4:
        return [1600, 1550, 1950, 4260]
    if cols == 5:
        return [1700, 1350, 900, 3260, 2150]
    base = PAGE_WIDTH_DXA // cols
    widths = [base] * cols
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc, rows):
    widths = choose_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        prevent_row_split(table.rows[row_index])
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            if row_index == 0:
                set_cell_shading(cell, HEADER_DARK)
            elif row_index % 2 == 0:
                set_cell_shading(cell, ROW_ALT_FILL)
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)
            add_inline(
                paragraph,
                value,
                size=8.8 if len(rows[0]) >= 4 else 9.2,
                color=WHITE if row_index == 0 else INK,
                bold=row_index == 0 or (len(rows[0]) == 2 and col_index == 0),
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    set_run_font(p.add_run("DSS301  |  Agricultural Drone Flight Scheduler"), size=9, color=MUTED, bold=True)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_title_block(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4, line=1.0)
    set_run_font(p.add_run("DSS301  |  WEEK 4 SUBMISSION"), size=9, color=BLUE, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4, line=1.0)
    set_run_font(p.add_run("DATA DICTIONARY"), size=27, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=10, line=1.0)
    set_run_font(p.add_run("DSS301 - Agricultural Drone Flight Scheduler"), size=14, color=BLUE, bold=True)

    metadata = [
        ("Đề tài", "DSS cho lập lịch bay UAV nông nghiệp theo thời tiết"),
        ("Nguồn dữ liệu", "WeatherAPI Forecast API (https://www.weatherapi.com)"),
        ("Cập nhật lần cuối", "2026-06-01"),
        ("Người phụ trách", "Data Engineer"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2100, 7260])
    set_table_borders(table)
    for row, (label, value) in zip(table.rows, metadata):
        prevent_row_split(row)
        set_cell_shading(row.cells[0], HEADER_FILL)
        set_cell_shading(row.cells[1], "FBFCFE")
        label_p = row.cells[0].paragraphs[0]
        set_paragraph_spacing(label_p, before=0, after=0, line=1.0)
        set_run_font(label_p.add_run(label.upper()), size=8.8, color=DARK_BLUE, bold=True)
        value_p = row.cells[1].paragraphs[0]
        set_paragraph_spacing(value_p, before=0, after=0, line=1.0)
        add_inline(value_p, value, size=9.4, color=INK)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=6, line=1.0)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    add_summary_strip(doc)


def parse_markdown(doc, lines):
    index = 8
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index].rstrip()

        if line.startswith("```"):
            if in_code:
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                set_table_geometry(table, [PAGE_WIDTH_DXA])
                set_table_borders(table)
                cell = table.cell(0, 0)
                set_cell_shading(cell, "F8FAFC")
                paragraph = cell.paragraphs[0]
                set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)
                set_run_font(paragraph.add_run("\n".join(code_lines)), name="Consolas", size=8.8, color=DARK_BLUE)
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[\s|:-]+\|$", lines[index + 1].strip()):
            rows = []
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_markdown_table(doc, rows)
            continue

        if not line or line == "---":
            index += 1
            continue
        if line.startswith("### "):
            add_body(doc, line[4:], style="Heading 2", after=7, line=1.0)
        elif line.startswith("## "):
            add_body(doc, line[3:], style="Heading 1", after=10, line=1.0)
        elif line.startswith("# "):
            add_body(doc, line[2:], style="Heading 1", after=10, line=1.0)
        elif line.startswith("> "):
            add_note(doc, line[2:])
        elif re.match(r"^\d+\. ", line):
            add_body(doc, re.sub(r"^\d+\. ", "", line), style="List Number", after=4)
        elif line.startswith("- "):
            add_body(doc, line[2:], style="List Bullet", after=4)
        else:
            add_body(doc, line)
        index += 1


def build():
    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8").splitlines())
    core = doc.core_properties
    core.title = "Data Dictionary - Agricultural Drone Flight Scheduler"
    core.subject = "DSS301 Week 4 submission"
    core.author = "DSS301 Project Team"
    core.keywords = "DSS301, WeatherAPI, clean dataset, agricultural drone"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
